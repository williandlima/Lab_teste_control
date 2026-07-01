"""Relatório FCT em Word (.docx) via python-docx, layout vindo de `report_template.yaml`.

Mesma identidade visual do relatório Excel (`excel_report.py`): cabeçalhos
em azul-marinho com texto claro, e o resultado avaliado pelo operador
destacado na cor semântica correspondente (verde/vermelho/amarelo), nunca
calculado — apenas refletindo o que já foi gravado em `evaluations`.
"""
from __future__ import annotations

import tempfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import _Cell

from config import BrandingConfig
from reports.chart import render_samples_chart
from reports.report_data import ReportData
from reports.template_engine import (
    build_context,
    evenly_sampled,
    load_template,
    render_fields,
    resolve_output_path,
    result_color_hex,
    step_stats_rows,
)

_LABEL_COL_WIDTH_IN = 2.2
_VALUE_COL_WIDTH_IN = 3.8


def _rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color.lstrip("#"))


def _shade_cell(cell: _Cell, hex_color: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    cell._tc.get_or_add_tcPr().append(shd)


def _add_heading(doc: Document, text: str, branding: BrandingConfig) -> None:
    heading = doc.add_heading(text, level=2)
    for run in heading.runs:
        run.font.color.rgb = _rgb(branding.color_secondary_navy)


def _fix_column_widths(table, widths_in: list[float]) -> None:
    """Larguras fixas de coluna (mesma proporção do PDF: 2,2"/3,8").

    Sem `table.autofit = False`, o Word ignora a largura pedida e reparte o
    espaço automaticamente — valores longos (ex.: resposta *IDN?) quebram de
    forma imprevisível. E a largura precisa ser setada em CADA célula, não só
    na coluna (peculiaridade do formato .docx: a coluna não é uma entidade
    própria, é inferida das larguras de célula da 1ª linha).
    """
    table.autofit = False
    for row in table.rows:
        for idx, width_in in enumerate(widths_in):
            row.cells[idx].width = Inches(width_in)


def _repeat_header_row(row) -> None:
    """Marca a linha como cabeçalho repetido em toda página (w:tblHeader).

    Sem isso, uma tabela de amostras com centenas de linhas quebra em várias
    páginas no Word e só a primeira mostra os nomes das colunas — o PDF já
    resolve isso via `repeatRows=1` do reportlab (ver `pdf_report.py`).
    """
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def _add_fields_table(doc: Document, fields: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in fields:
        row = table.add_row()
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        row.cells[1].text = value
    _fix_column_widths(table, [_LABEL_COL_WIDTH_IN, _VALUE_COL_WIDTH_IN])


def _add_data_table(
    doc: Document, columns: list[str], rows: list[list[str]], branding: BrandingConfig
) -> None:
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for idx, header in enumerate(columns):
        header_cells[idx].text = header
        _shade_cell(header_cells[idx], branding.color_primary_navy)
        run = header_cells[idx].paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = _rgb(branding.color_text_on_navy)
    _repeat_header_row(table.rows[0])
    for data_row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(data_row):
            cells[idx].text = value


def _add_page_number_field(paragraph, field_code: str) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {field_code} "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def _add_page_number_footer(doc: Document) -> None:
    """'Página X de Y' no rodapé de página (campos PAGE/NUMPAGES do Word).

    Necessário para reimpressão/arquivamento de relatórios longos — sem isso
    não dá pra saber se uma página impressa avulsa pertence a este relatório
    nem se faltam páginas.
    """
    footer_paragraph = doc.sections[0].footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_paragraph.add_run("Página ")
    _add_page_number_field(footer_paragraph, "PAGE")
    footer_paragraph.add_run(" de ")
    _add_page_number_field(footer_paragraph, "NUMPAGES")
    for run in footer_paragraph.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


def generate_word_report(
    data: ReportData, branding: BrandingConfig, output_dir: Path, base_name: str | None = None
) -> Path:
    template = load_template()
    context = build_context(data, branding)

    doc = Document()
    _add_page_number_footer(doc)

    if branding.logo_path.exists():
        doc.add_picture(str(branding.logo_path), width=Inches(1.0))

    title = doc.add_heading(template["title"], level=1)
    for run in title.runs:
        run.font.color.rgb = _rgb(branding.color_primary_navy)

    subtitle = doc.add_paragraph(context["company_name"])
    subtitle.runs[0].font.italic = True
    subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for key in ("identification", "parameters", "execution", "traceability"):
        section = template["sections"][key]
        _add_heading(doc, section["heading"], branding)
        _add_fields_table(doc, render_fields(section["fields"], context))

    evaluation_section = template["sections"]["evaluation"]
    _add_heading(doc, evaluation_section["heading"], branding)
    eval_fields = render_fields(evaluation_section["fields"], context)
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    color = result_color_hex(data, template, branding)
    for idx, (label, value) in enumerate(eval_fields):
        row = table.add_row()
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        row.cells[1].text = value
        if idx == 0 and color:
            _shade_cell(row.cells[1], color)
    _fix_column_widths(table, [_LABEL_COL_WIDTH_IN, _VALUE_COL_WIDTH_IN])

    chart_png: Path | None = None
    chart_section = template["sections"]["chart"]
    with tempfile.TemporaryDirectory() as tmp_dir:
        chart_png = render_samples_chart(data, branding, Path(tmp_dir) / "chart.png")
        if chart_png is not None:
            _add_heading(doc, chart_section["heading"], branding)
            doc.add_picture(str(chart_png), width=Inches(6.2))

        stats_section = template["sections"]["step_stats_table"]
        stats_rows = step_stats_rows(data)
        if stats_rows:
            _add_heading(doc, stats_section["heading"], branding)
            _add_data_table(doc, stats_section["columns"], stats_rows, branding)

        samples_section = template["sections"]["samples_table"]
        _add_heading(doc, samples_section["heading"], branding)
        sampled = evenly_sampled(data.samples, samples_section["max_rows"])
        sample_rows = [
            [s.timestamp, str(s.step_index + 1), f"{s.current_measured:.3f}", f"{s.voltage_measured:.3f}"]
            for s in sampled
        ]
        _add_data_table(doc, samples_section["columns"], sample_rows, branding)

        events_section = template["sections"]["events_table"]
        _add_heading(doc, events_section["heading"], branding)
        event_rows = [[e.timestamp or "", e.level, e.source, e.message] for e in data.events]
        _add_data_table(doc, events_section["columns"], event_rows, branding)

        footer = doc.add_paragraph(template["footer"].format(**context))
        footer.runs[0].font.size = Pt(8)
        footer.runs[0].font.italic = True

        output_path = resolve_output_path(data, output_dir, "docx", base_name)
        doc.save(output_path)
    return output_path
